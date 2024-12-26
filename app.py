from flask import Flask, render_template, request, jsonify, send_file
import os
import tempfile
import re
import webbrowser
import openai
from openai import OpenAI
from docx import Document
import subprocess
import platform
from datetime import datetime
import json
import sys
from waitress import serve
from concurrent.futures import ThreadPoolExecutor

# 获取资源路径
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

app = Flask(__name__)

# 获取用户数据目录
if getattr(sys, 'frozen', False):
    user_data_dir = os.path.join(os.environ['APPDATA'], 'YourAppName') if os.name == 'nt' else os.path.expanduser('~/.yourappname')
else:
    user_data_dir = os.path.dirname(os.path.abspath(__file__))

os.makedirs(user_data_dir, exist_ok=True)

# 初始化 settings 变量
SETTINGS_FILE = os.path.join(user_data_dir, 'settings.json')
PROMPTS_FILE = os.path.join(user_data_dir, 'prompts.json')

if not os.path.exists(SETTINGS_FILE):
    default_settings = {
        "api_key": "",
        "model_version": "",
        "base_url": ""
    }
    with open(SETTINGS_FILE, 'w') as f:
        json.dump(default_settings, f)
else:
    with open(SETTINGS_FILE, 'r') as f:
        settings = json.load(f)

if not os.path.exists(PROMPTS_FILE):
    default_prompts = {
        "ctb_prompt": """我会给你短剧剧集的台词，你根据剧集台词写剧集的细纲，要求把剧集中剧情的转折和高潮细致剧情都反应出来。不分段 在一段中描述剧情。除了剧情之外不要有额外的分析。更简明的概括关键剧情。不要遗漏情绪的变化的剧情。
格式参考：

Emily Harper 是一位农场女工，在乡村种土豆和养鸡，并照顾体弱的母亲 Margaret Harper。一天早晨，她从朋友 Sarah Price 的手机上看到新闻，得知丈夫 Ethan Blake 作为 Northbrook Tech 的 CEO 即将主持公司 IPO 的敲钟仪式。Sarah 和 Margaret Turner 对此兴奋不已，但 Emily 的笑容中带着一丝失落。这时，秘书 Vanessa Price 突然驾车来到农场，带来一份 Ethan 委托的离婚协议。
""",
        "btc_prompt": """我会发你中文的单集细纲。请用好莱坞标准格式写作该剧本，地道的英语写作。每一集时长1-2分钟。
请用标准的好莱坞格式，粘贴到final draft软件中可直接识别
剧本中使用的英语地道，口语。
尽可能保留原有内容，对话和语气更贴近美剧台词，场景和动作描述细致
尽可能的丰富，可在对话中增加更多的冲突和肥皂剧情节
多运用英美剧中有的日常用俚语 口头语梗 玩笑梗 俏皮话
增加一些美剧中的语言风格，在保留原意的基础上改写台词。
理解意思之后用更地道的英语大幅度改写台词，不需要直译，可适当增加一些欧美梗和口头语
动作描述语言精简，只翻译结果动作；语言简洁，尽量通过动作传达情绪，而非形容词副词。
不要数字编号
仅写作剧本部分，不要做任何的额外分析
不要字体格式，比如集号，地点处
不要给Episode起名，Episode+集数即可，后面我会告诉你这是第几集，比如第一集就是 EPISODE 1 这段字符在final draft中以shot格式出现
人名只在整个剧本第一次出现时，所有字母大写，并后面括号标注年龄数字，数字后不加s。之后人名再次出现时，首字母大写，其余字母小写。
之前出现过的人名之后再次出现就不用大写和标注年龄了。
不用斜体，加粗等格式
集号放在正文里 不用粗体
不要fade in/ fade out / cut to这几个提示词
语气描述的parenthetical部分需要另起一行，并用小括号括起来
动作描述不要用括号。用标准的好莱坞格式。
不要出现△,*等特殊符号
一定要全英文，不要出现中文标点符号。
"""
    }
    with open(PROMPTS_FILE, 'w') as f:
        json.dump(default_prompts, f)
else:
    with open(PROMPTS_FILE, 'r') as f:
        prompts = json.load(f)

# 获取当前脚本所在的目录
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# 定义文件夹路径
FOLDERS = {
    'ctb_input': os.path.join(BASE_DIR, 'ctb_input'),
    'ctb_output': os.path.join(BASE_DIR, 'ctb_output'),
    'btc_input': os.path.join(BASE_DIR, 'btc_input'),
    'btc_output': os.path.join(BASE_DIR, 'btc_output')
}

# 确保所有文件夹都存在
for folder in FOLDERS.values():
    os.makedirs(folder, exist_ok=True)


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/update_settings', methods=['POST'])
def update_settings():
    global settings
    data = request.json
    settings['api_key'] = data.get('api_key', '')
    settings['model_version'] = data.get('model_version', '')
    settings['base_url'] = data.get('base_url', '')

    # 将设置保存到 JSON 文件
    with open(SETTINGS_FILE, 'w') as f:
        json.dump(settings, f)

    return jsonify({"success": True, "message": "设置已更新"})

@app.route('/load_settings', methods=['GET'])
def load_settings():
    if os.path.exists(SETTINGS_FILE):
        with open(SETTINGS_FILE, 'r') as f:
            settings = json.load(f)
    else:
        settings = {
            "api_key": "",
            "model_version": "",
            "base_url": ""
        }
    return jsonify(settings)

@app.route('/get_initial_folders', methods=['GET'])
def get_initial_folders():
    return jsonify(FOLDERS)

@app.route('/update_folders', methods=['POST'])
def update_folders():
    global FOLDERS
    data = request.json
    for key, value in data.items():
        if key in FOLDERS:
            FOLDERS[key] = value
            os.makedirs(value, exist_ok=True)
    return jsonify(success=True)

@app.route('/get_default_prompts', methods=['GET'])
def get_default_prompts():
    if os.path.exists(PROMPTS_FILE):
        with open(PROMPTS_FILE, 'r') as f:
            prompts = json.load(f)
    else:
        prompts = {
            "ctb_prompt": "",
            "btc_prompt": ""
        }
    return jsonify(prompts)

@app.route('/update_prompts', methods=['POST'])
def update_prompts():
    global prompts
    data = request.json
    prompts['ctb_prompt'] = data.get('ctb_prompt', '')
    prompts['btc_prompt'] = data.get('btc_prompt', '')

    # 将提示词保存到 JSON 文件
    with open(PROMPTS_FILE, 'w') as f:
        json.dump(prompts, f)

    return jsonify({"success": True, "message": "提示词已更新"})

# 全局变量来存储进度信息
progress_info = {}

def process_episode(api_key,base_url, prompt, episode_content, episode_index):
    full_prompt = f"{prompt}\n这是第{episode_index}集\n{episode_content}"
    client = openai.Client(api_key=api_key,base_url=base_url)
    try:
        response = client.chat.completions.create(
            model=settings.get('model_version'),
            messages=[{"role": "user", "content": full_prompt}]
        )
        return response.choices[0].message.content, None
    except Exception as e:
        return None, str(e)

@app.route('/ctb_process', methods=['POST'])
def ctb_process():
    api_key = request.form['api_key']
    prompt = request.form['prompt']
    base_url = request.form['base_url']
    
    openai.api_key = api_key
    openai.base_url = base_url
    # 创建临时文件夹来存储上传的文件
    with tempfile.TemporaryDirectory() as temp_dir:
        input_folder = temp_dir
        output_folder = FOLDERS['ctb_output']
        
        # 获取当前时间戳
        now = datetime.now()
        date_str = now.strftime("%Y%m%d")  # 年月日
        timestamp_str = now.strftime("%Y%m%d%H%M%S%f")  # 年月日时分秒微秒
        
        # 创建新的文件夹
        new_output_folder = os.path.join(output_folder, date_str)
        if not os.path.exists(new_output_folder):
            os.makedirs(new_output_folder)
        
        result_doc = Document()
        files = request.files.getlist('files[]')
        total_episodes = 0
        processed_episodes = 0
        
        for file in files:
            file_path = os.path.join(input_folder, file.filename)
            file.save(file_path)
            print(f"File saved to: {file_path}")  # 打印文件路径以验证
            
            if not os.path.exists(file_path):
                return jsonify(success=False, error=f"文件未找到: {file_path}"), 400
            
            doc = Document(file_path)
            content = '\n'.join([para.text for para in doc.paragraphs])
            
            # 使用正则表达式来分割内容，匹配 "第*集\n"
            episodes = re.split(r'(第\d+集\n)', content)
            total_episodes += len(episodes) // 2
        
        progress_info[timestamp_str] = {'current_episode': 0}
        
        def update_progress(current_episode):
            progress_info[timestamp_str]['current_episode'] = current_episode
        
        with ThreadPoolExecutor(max_workers=2) as executor:
            futures = []
            current_episode_index = 1
            
            for file in files:
                file_path = os.path.join(input_folder, file.filename)
                
                if not os.path.exists(file_path):
                    return jsonify(success=False, error=f"文件未找到: {file_path}"), 400
                
                doc = Document(file_path)
                content = '\n'.join([para.text for para in doc.paragraphs])
                
                # 使用正则表达式来分割内容，匹配 "第*集\n"
                episodes = re.split(r'(第\d+集\n)', content)
                
                for j, episode in enumerate(episodes):
                    if re.match(r'第(\d+)集\n', episode):
                        # 提取集数
                        match = re.match(r'第(\d+)集\n', episode)
                        episode_number = int(match.group(1))
                        current_episode_index = episode_number
                    elif episode.strip():  # 确保段落不为空
                        future = executor.submit(process_episode, api_key, base_url, prompt, episode, current_episode_index)
                        futures.append((future, current_episode_index))
                        current_episode_index += 1
            
            for future, episode_index in futures:
                try:
                    response_content, error = future.result()
                    if response_content is None:
                        print(f"Error processing episode {episode_index}: {error}")
                    else:
                        result_doc.add_heading(f'EPISODE {episode_index}', level=1)
                        result_doc.add_paragraph(response_content)
                        update_progress(episode_index)
                except Exception as e:
                    print(f"Exception occurred while processing episode {episode_index}: {str(e)}")
        
        # 使用时间戳命名文件
        name_time = now.strftime("%Y%m%d%H%M")
        output_file = os.path.join(new_output_folder, f"输出文件_{name_time}.docx")
        result_doc.save(output_file)
    
    del progress_info[timestamp_str]
    return jsonify(success=True, message="台词转台本处理完成", output_file=output_file)


def process_paragraph(api_key,base_url, prompt, paragraph_content, episode_index):
    full_prompt = f"{prompt}\n这是第{episode_index}集\n{paragraph_content}"
    print(f"这是第{episode_index}集")
    client = openai.Client(api_key=api_key,base_url=base_url)
    try:
        response = client.chat.completions.create(
            model=settings.get('model_version'),
            messages=[{"role": "user", "content": full_prompt}]
        )
        return response.choices[0].message.content, None
    except Exception as e:
        return None, str(e)

@app.route('/btc_process', methods=['POST'])
def btc_process():
    api_key = request.form['api_key']
    prompt = request.form['prompt']
    base_url = request.form['base_url']
    openai.base_url = base_url
    openai.api_key = api_key
    
    # 创建临时文件夹来存储上传的文件
    with tempfile.TemporaryDirectory() as temp_dir:
        input_folder = temp_dir
        output_folder = FOLDERS['btc_output']
        
        # 获取当前时间戳
        now = datetime.now()
        timestamp_str = now.strftime("%Y%m%d%H%M%S%f")  # 年月日时分秒微秒
        
        # 创建新的文件夹
        new_output_folder = os.path.join(output_folder, now.strftime("%Y%m%d%H%M"))
        if not os.path.exists(new_output_folder):
            os.makedirs(new_output_folder)
        
        file = request.files['file']
        file_path = os.path.join(input_folder, file.filename)
        file.save(file_path)
        print(f"File saved to: {file_path}")  # 打印文件路径以验证
        
        if not os.path.exists(file_path):
            return jsonify(success=False, error=f"文件未找到: {file_path}"), 400
        
        doc = Document(file_path)
        text_content = '\n'.join([para.text for para in doc.paragraphs])
        
        # 使用正则表达式按“第几集”分割文本
        pattern = r'(?s)(?:第([零一二三四五六七八九十百千万\d]+)集|EPISODE\s+([零一二三四五六七八九十百千万\d]+))\s*(.*?)(?=第([零一二三四五六七八九十百千万\d]+)集|EPISODE\s+[零一二三四五六七八九十百千万\d]+|\Z)'
        matches = list(re.finditer(pattern, text_content))

        # 提取实际内容部分
        episodes = []
        for match in matches:
            episode_number_1 = match.group(1)
            episode_number_2 = match.group(2)
            content = match.group(3).strip()
            
            if episode_number_2 is None and episode_number_1 is not None:
                episodes.append(content)
            elif episode_number_2 is not None:
                episodes.append(content)

        print(episodes)
        # 调试信息：打印提取的内容
        print("Extracted episodes:", episodes)
        
        result_doc = Document()
        total_episodes = len(episodes)
        processed_episodes = 0
        
        progress_info[timestamp_str] = {'current_episode': 0, 'total_episodes': total_episodes}
        
        def update_progress(current_episode):
            progress_info[timestamp_str]['current_episode'] = current_episode
        
        def retry_process_paragraph(api_key, base_url, prompt, paragraph_content, episode_index, max_retries=3):
            attempt = 0
            while attempt < max_retries:
                response_content, error = process_paragraph(api_key, base_url, prompt, paragraph_content, episode_index)
                if response_content is not None:
                    return response_content, None
                elif "504" in error:
                    attempt += 1
                    print(f"Retrying episode {episode_index}, attempt {attempt}")
                else:
                    return None, error
            return None, "Max retries exceeded"
        
        with ThreadPoolExecutor(max_workers=5) as executor:  # 增加并发线程数为5
            futures = []
            current_episode_index = 1
            
            for paragraph in episodes:
                future = executor.submit(retry_process_paragraph, api_key, base_url, prompt, paragraph.strip(), current_episode_index)
                print(f"Submitting task for episode {current_episode_index}")
                futures.append((future, current_episode_index))
                current_episode_index += 1
            
            for future, episode_index in futures:
                try:
                    response_content, error = future.result()
                    if response_content is None:
                        print(f"Error processing episode {episode_index}: {error}")
                    else:
                        result_doc.add_paragraph(response_content)
                        update_progress(episode_index)
                        print(f"Processed episode {episode_index}")
                except Exception as e:
                    print(f"Exception occurred while processing episode {episode_index}: {str(e)}")
        
        # 使用时间戳命名文件
        output_file = os.path.join(new_output_folder, f'combined_episodes.docx')
        result_doc.save(output_file)
    
    del progress_info[timestamp_str]
    return jsonify(success=True, message="台本转台词处理完成", output_file=output_file)

@app.route('/get_progress/<timestamp>', methods=['GET'])
def get_progress(timestamp):
    if timestamp in progress_info:
        return jsonify(progress_info[timestamp])
    else:
        return jsonify({'current_episode': 0})




@app.route('/open_folder/<folder_name>', methods=['GET'])
def open_folder(folder_name):
    if folder_name not in FOLDERS:
        return jsonify({"success": False, "message": "文件夹不存在"}), 404
    
    folder_path = FOLDERS[folder_name]
    
    try:
        if platform.system() == "Windows":
            os.startfile(folder_path)
        elif platform.system() == "Darwin":  # macOS
            subprocess.Popen(["open", folder_path])
        elif platform.system() == "Linux":
            subprocess.Popen(["xdg-open", folder_path])
        else:
            return jsonify({"success": False, "message": "操作系统不受支持"}), 501
        
        return jsonify({"success": True, "message": "文件夹已打开"})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500
@app.route('/shutdown', methods=['POST'])
def shutdown():
    try:
        # 强制退出进程
        os._exit(0)
        return jsonify(success=True, message="程序已关闭"), 200
    except Exception as e:
        return jsonify(success=False, error=str(e)), 500

if __name__ == '__main__':
    webbrowser.open('http://127.0.0.1:8091')
    # app.run(debug=True, host='0.0.0.0',port=8091,use_reloader=False)
    serve(app, host='0.0.0.0', port=8091)


